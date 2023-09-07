import pandas as pd
import datetime
from docx import Document

def generate_dates(start_date, end_date):
    start_dates = []
    end_dates = []
    day_diff = []
    while start_date <= end_date:
        start_formatted_date = start_date.strftime('%m/%d/%Y')
        end_formatted_date = (start_date + pd.DateOffset(months=3) - pd.DateOffset(days=1)).strftime('%m/%d/%Y')
        diff = (start_date + pd.DateOffset(months=3) - pd.DateOffset(days=1) - start_date).days + 1

        start_dates.append(start_formatted_date)
        end_dates.append(end_formatted_date)
        day_diff.append(diff)

        year = start_date.year
        month = start_date.month

        if month == 10:
            year += 1
            month = 1
        else:
            month = (month + 3) % 12

        start_date = datetime.datetime(year=year, month=month, day=1)

    if start_dates[-1] != end_date.strftime('%m/%d/%Y'):
        start_dates.append(start_date.strftime('%m/%d/%Y'))
        end_dates.append(end_date.strftime('%m/%d/%Y'))
        day_diff.append((end_date - start_date).days + 1)

    return start_dates,end_dates,day_diff

def EOQ (date, dfdata):
    return (dfdata.loc[date, "Ending Dates"])

def days_in_quarter (date, dfdata):
    return (dfdata.loc[date, 'Days']) 

def get_first_day_of_next_month(dt):
    year = dt.year
    month = dt.month
    if month == 12:
        year += 1
        month = 1
    else:
        month += 1
    first_day_of_next_month = datetime.datetime(year=year, month=month, day=1)
    return first_day_of_next_month

def main_idea (df, dfdata, Savings_Rate, Final_Val_Days):
    for m in range(len(df)):
        VTD_Date = df["VTD"][m]
        # Added Code
        if VTD_Date > dfdata["Ending Dates"][len(dfdata)-1]:
            new_row_starting_date = get_first_day_of_next_month(dfdata["Ending Dates"][len(dfdata)-1])
            new_row_days = (VTD_Date - new_row_starting_date).days
            new_row_days = new_row_days + 1
            new_row_returns = (float(new_row_days)*(float(Savings_Rate.strip("%"))))/(float(365))
            new_row = {'Starting Dates': new_row_starting_date, 
                    'Ending Dates': VTD_Date, 
                    'Days': new_row_days,
                    'Returns': "{:.2f}%".format(new_row_returns)}
            new_row_df = pd.DataFrame([new_row])
            dfdata = pd.concat([dfdata, new_row_df], ignore_index=True)
        # End of Added Code
        Pay_Date = df["Pay Date"][m]
        mask1 = (dfdata['Starting Dates'] <= Pay_Date) & (dfdata['Ending Dates'] >= Pay_Date)
        Loop_Beginning = dfdata.loc[mask1].index[0]
        mask2 = (dfdata['Starting Dates'] <= VTD_Date) & (dfdata['Ending Dates'] >= VTD_Date)
        Loop_Ending = dfdata.loc[mask2].index[0]
        month_first_day = Pay_Date.replace(day=1)
        # if x > 0, j is days considered during quarter for when there is positive accumulation  
        j = (EOQ(Loop_Beginning, dfdata) - month_first_day).days
        j = j+1
        first_day_next_month = get_first_day_of_next_month(Pay_Date)
        # if x < 0, k is days considered during quarter for when there is negative accumulation
        k = (EOQ(Loop_Beginning, dfdata) - first_day_next_month).days
        k = k+1
        x = float(dfdata.loc[Loop_Beginning, 'Returns'].replace('%', ''))
        if x > 0:
            y = ((df["Corrective Contribution"][m])*(1+(x/100))**(j/(days_in_quarter(Loop_Beginning, dfdata))))
        else:
            y = ((df["Corrective Contribution"][m])*(1+(x/100))**(k/(days_in_quarter(Loop_Beginning, dfdata))))

        calc_df = dfdata.copy()
        calc_df["Compounding"] = 0
        calc_df["Compounding"][Loop_Beginning] = y
        calc_df["Total Compounding"] = 0
        if Loop_Beginning == Loop_Ending:
            calc_df["Total Compounding"] = calc_df["Compounding"]
        else:
            calc_df["Total Compounding"][Loop_Beginning] = calc_df["Compounding"][Loop_Beginning]
            for i in range(Loop_Beginning+1, Loop_Ending+1):
                calc_df["Total Compounding"][i] = calc_df["Total Compounding"][i-1] \
                *(1+(float(calc_df["Returns"][i].replace("%",""))/100))

        datedif = str(EOQ(Loop_Ending, dfdata) - VTD_Date)
        datedif = float(datedif[:2])
        calc_df["TMV VTD"] = 0
        calc_df["TMV VTD"][Loop_Ending] = calc_df["Total Compounding"][Loop_Ending]* \
        (1+(float(calc_df["Returns"][Loop_Ending].replace("%",""))/100))**(-datedif/days_in_quarter(Loop_Ending, dfdata))
        calc_df["Net VTD"] = 0
        calc_df["Net VTD"][Loop_Ending] = calc_df["TMV VTD"][Loop_Ending] - df["Corrective Contribution"][m]
        calc_df["Net to EOQ"] = calc_df["Net VTD"]* \
        (1+(float(calc_df["Returns"][Loop_Ending].replace("%",""))/100))**(datedif/days_in_quarter(Loop_Ending, dfdata))
        calc_df["Bring Returns Forward"] = 0
        calc_df["Bring Returns Forward"][Loop_Ending] = calc_df["Net to EOQ"][Loop_Ending]
        for i in range(Loop_Ending+1, len(calc_df)):
            calc_df["Bring Returns Forward"][i] = calc_df["Bring Returns Forward"][i-1] \
            *(1+(float(calc_df["Returns"][i].replace("%",""))/100))
        calc_df["At Present"] = 0
        calc_df["At Present"][len(calc_df)-1] = calc_df["Bring Returns Forward"][len(calc_df)-1] \
        *(1+(float(Savings_Rate.replace("%",""))/100))**(Final_Val_Days/365)
        answer = round(calc_df["At Present"][len(calc_df)-1],2)
        df["Investment Earnings"][m] = answer
        if VTD_Date > dfdata["Ending Dates"][len(dfdata)-1]:
            dfdata = dfdata.drop(dfdata.index[-1])
    return calc_df

def positive_procedure (df, Single_VTD):
    df_positive = df[:-1]
    df_positive = df_positive[df_positive["Investment Earnings"] > 0]
    numeric_columns = ['Corrective Contribution', 'Investment Earnings']
    df_positive = df_positive.groupby(['Name'], as_index=False)[numeric_columns].sum(numeric_only=True)
    df_positive['Name'] = df_positive['Name'].astype(str)
    df_positive['Year'] = df_positive['Name'].str.split().str[-1]
    df_positive['Year'] = pd.to_numeric(df_positive['Year'])
    df_positive = df_positive.sort_values(by=['Year', 'Name'])
    df_positive.drop('Year', axis=1, inplace=True)
    sum_row = df_positive.sum(numeric_only=True)
    sum_df = pd.DataFrame([sum_row], columns=df.columns)
    df_positive = pd.concat([df_positive, sum_df], ignore_index=True)
    df_positive = df_positive[["Name", "Corrective Contribution", "Investment Earnings"]]
    df_positive.loc[len(df_positive)-1, "Name"] = "Total"
    output_file_name_1 = f"Positive Earnings VTD {Single_VTD}.xlsx"
    df_positive.to_excel(output_file_name_1, index=False)
    print(f"Dataframe exported to {output_file_name_1}")
    return df_positive

def negative_procedure (df, Single_VTD):
    df_negative = df[:-1]
    df_negative = df_negative[df_negative["Investment Earnings"] < 0]
    numeric_columns = ['Corrective Contribution', 'Investment Earnings']
    df_negative = df_negative.groupby(['Name'], as_index=False)[numeric_columns].sum(numeric_only=True)
    df_negative['Name'] = df_negative['Name'].astype(str)
    df_negative['Year'] = df_negative['Name'].str.split().str[-1]
    df_negative['Year'] = pd.to_numeric(df_negative['Year'])
    df_negative = df_negative.sort_values(by=['Year', 'Name'])
    df_negative.drop('Year', axis=1, inplace=True)
    sum_row = df_negative.sum(numeric_only=True)
    sum_df = pd.DataFrame([sum_row], columns=df.columns)
    df_negative = pd.concat([df_negative, sum_df], ignore_index=True)
    df_negative = df_negative[["Name", "Corrective Contribution", "Investment Earnings"]]
    df_negative['"Recommended" Investment Earnings'] = 0
    df_negative.loc[len(df_negative)-1, "Name"] = "Total"
    output_file_name_2 = f"Negative Earnings VTD {Single_VTD}.xlsx"
    df_negative.to_excel(output_file_name_2, index=False)
    print(f"Dataframe exported to {output_file_name_2}")  
    return df_negative

def summary_procedure (df, Single_VTD, contains_negative):
    df_summary = df[["Name", "Corrective Contribution", "Investment Earnings"]]
    numeric_columns = ['Corrective Contribution', 'Investment Earnings']
    df_summary = df_summary.groupby(['Name'], as_index=False)[numeric_columns].sum(numeric_only=True)
    df_summary['Name'] = df_summary['Name'].astype(str)
    df_summary['Year'] = df_summary['Name'].str.split().str[-1]
    df_summary['Year'] = pd.to_numeric(df_summary['Year'])
    df_summary = df_summary.sort_values(by=['Year', 'Name'])
    df_summary.drop('Year', axis=1, inplace=True)
    sum_row = df_summary.sum(numeric_only=True)
    sum_df = pd.DataFrame([sum_row], columns=df.columns)
    df_summary = pd.concat([df_summary, sum_df], ignore_index=True)
    df_summary = df_summary[["Name", "Corrective Contribution", "Investment Earnings"]]
    if contains_negative == "yes":
        df_summary['"Recommended" Investment Earnings'] = 0
    df_summary.loc[len(df_summary)-1, "Name"] = "Total"
    output_file_name = f"Summary VTD {Single_VTD}.xlsx"
    df_summary.to_excel(output_file_name, index=False)
    print(f"Dataframe exported to {output_file_name}")
    return df_summary

def main_idea_2 (df, dfdata, Savings_Rate, Final_Val_Days):
    for m in range(len(df)):
        VTD_Date = df["VTD"][m]
        # Added Code
        if VTD_Date > dfdata["Ending Dates"][len(dfdata)-1]:
            new_row_starting_date = get_first_day_of_next_month(dfdata["Ending Dates"][len(dfdata)-1])
            new_row_days = (VTD_Date - new_row_starting_date).days
            new_row_days = new_row_days + 1
            new_row_returns = (float(new_row_days)*(float(Savings_Rate.strip("%"))))/(float(365))
            new_row = {'Starting Dates': new_row_starting_date, 
                    'Ending Dates': VTD_Date, 
                    'Days': new_row_days,
                    'Returns': "{:.2f}%".format(new_row_returns)}
            new_row_df = pd.DataFrame([new_row])
            dfdata = pd.concat([dfdata, new_row_df], ignore_index=True)
        # End of Added Code
        Pay_Date = df["Pay Date"][m]
        mask1 = (dfdata['Starting Dates'] <= Pay_Date) & (dfdata['Ending Dates'] >= Pay_Date)
        Loop_Beginning = dfdata.loc[mask1].index[0]
        mask2 = (dfdata['Starting Dates'] <= VTD_Date) & (dfdata['Ending Dates'] >= VTD_Date)
        Loop_Ending = dfdata.loc[mask2].index[0]
        month_first_day = Pay_Date.replace(day=1)
        # if x > 0, j is days considered during quarter for when there is positive accumulation  
        j = (EOQ(Loop_Beginning, dfdata) - month_first_day).days
        j = j+1
        first_day_next_month = get_first_day_of_next_month(Pay_Date)
        # if x < 0, k is days considered during quarter for when there is negative accumulation
        k = (EOQ(Loop_Beginning, dfdata) - first_day_next_month).days
        k = k+1
        x = float(dfdata.loc[Loop_Beginning, 'Returns'].replace('%', ''))
        if x > 0:
            y = ((df["EE Corrective Contribution"][m])*(1+(x/100))**(j/(days_in_quarter(Loop_Beginning, dfdata))))
        else:
            y = ((df["EE Corrective Contribution"][m])*(1+(x/100))**(k/(days_in_quarter(Loop_Beginning, dfdata))))

        calc_df = dfdata.copy()
        calc_df["Compounding"] = 0
        calc_df["Compounding"][Loop_Beginning] = y
        calc_df["Total Compounding"] = 0
        if Loop_Beginning == Loop_Ending:
            calc_df["Total Compounding"] = calc_df["Compounding"]
        else:
            calc_df["Total Compounding"][Loop_Beginning] = calc_df["Compounding"][Loop_Beginning]
            for i in range(Loop_Beginning+1, Loop_Ending+1):
                calc_df["Total Compounding"][i] = calc_df["Total Compounding"][i-1] \
                *(1+(float(calc_df["Returns"][i].replace("%",""))/100))

        datedif = str(EOQ(Loop_Ending, dfdata) - VTD_Date)
        datedif = float(datedif[:2])
        calc_df["TMV VTD"] = 0
        calc_df["TMV VTD"][Loop_Ending] = calc_df["Total Compounding"][Loop_Ending]* \
        (1+(float(calc_df["Returns"][Loop_Ending].replace("%",""))/100))**(-datedif/days_in_quarter(Loop_Ending, dfdata))
        calc_df["Net VTD"] = 0
        calc_df["Net VTD"][Loop_Ending] = calc_df["TMV VTD"][Loop_Ending] - df["EE Corrective Contribution"][m]
        calc_df["Net to EOQ"] = calc_df["Net VTD"]* \
        (1+(float(calc_df["Returns"][Loop_Ending].replace("%",""))/100))**(datedif/days_in_quarter(Loop_Ending, dfdata))
        calc_df["Bring Returns Forward"] = 0
        calc_df["Bring Returns Forward"][Loop_Ending] = calc_df["Net to EOQ"][Loop_Ending]
        for i in range(Loop_Ending+1, len(calc_df)):
            calc_df["Bring Returns Forward"][i] = calc_df["Bring Returns Forward"][i-1] \
            *(1+(float(calc_df["Returns"][i].replace("%",""))/100))
        calc_df["At Present"] = 0
        calc_df["At Present"][len(calc_df)-1] = calc_df["Bring Returns Forward"][len(calc_df)-1] \
        *(1+(float(Savings_Rate.replace("%",""))/100))**(Final_Val_Days/365)
        answer = round(calc_df["At Present"][len(calc_df)-1],2)
        df["EE Investment Earnings"][m] = answer
        if VTD_Date > dfdata["Ending Dates"][len(dfdata)-1]:
            dfdata = dfdata.drop(dfdata.index[-1])
    #########################################################
    for m in range(len(df)):
        VTD_Date = df["VTD"][m]
        # Added Code
        if VTD_Date > dfdata["Ending Dates"][len(dfdata)-1]:
            new_row_starting_date = get_first_day_of_next_month(dfdata["Ending Dates"][len(dfdata)-1])
            new_row_days = (VTD_Date - new_row_starting_date).days
            new_row_days = new_row_days + 1
            new_row_returns = (float(new_row_days)*(float(Savings_Rate.strip("%"))))/(float(365))
            new_row = {'Starting Dates': new_row_starting_date, 
                    'Ending Dates': VTD_Date, 
                    'Days': new_row_days,
                    'Returns': "{:.2f}%".format(new_row_returns)}
            new_row_df = pd.DataFrame([new_row])
            dfdata = pd.concat([dfdata, new_row_df], ignore_index=True)
        # End of Added Code
        Pay_Date = df["Pay Date"][m]
        mask1 = (dfdata['Starting Dates'] <= Pay_Date) & (dfdata['Ending Dates'] >= Pay_Date)
        Loop_Beginning = dfdata.loc[mask1].index[0]
        mask2 = (dfdata['Starting Dates'] <= VTD_Date) & (dfdata['Ending Dates'] >= VTD_Date)
        Loop_Ending = dfdata.loc[mask2].index[0]
        month_first_day = Pay_Date.replace(day=1)
        # if x > 0, j is days considered during quarter for when there is positive accumulation  
        j = (EOQ(Loop_Beginning, dfdata) - month_first_day).days
        j = j+1
        first_day_next_month = get_first_day_of_next_month(Pay_Date)
        # if x < 0, k is days considered during quarter for when there is negative accumulation
        k = (EOQ(Loop_Beginning, dfdata) - first_day_next_month).days
        k = k+1
        x = float(dfdata.loc[Loop_Beginning, 'Returns'].replace('%', ''))
        if x > 0:
            y = ((df["ER Corrective Contribution"][m])*(1+(x/100))**(j/(days_in_quarter(Loop_Beginning, dfdata))))
        else:
            y = ((df["ER Corrective Contribution"][m])*(1+(x/100))**(k/(days_in_quarter(Loop_Beginning, dfdata))))

        calc_df = dfdata.copy()
        calc_df["Compounding"] = 0
        calc_df["Compounding"][Loop_Beginning] = y
        calc_df["Total Compounding"] = 0
        if Loop_Beginning == Loop_Ending:
            calc_df["Total Compounding"] = calc_df["Compounding"]
        else:
            calc_df["Total Compounding"][Loop_Beginning] = calc_df["Compounding"][Loop_Beginning]
            for i in range(Loop_Beginning+1, Loop_Ending+1):
                calc_df["Total Compounding"][i] = calc_df["Total Compounding"][i-1] \
                *(1+(float(calc_df["Returns"][i].replace("%",""))/100))

        datedif = str(EOQ(Loop_Ending, dfdata) - VTD_Date)
        datedif = float(datedif[:2])
        calc_df["TMV VTD"] = 0
        calc_df["TMV VTD"][Loop_Ending] = calc_df["Total Compounding"][Loop_Ending]* \
        (1+(float(calc_df["Returns"][Loop_Ending].replace("%",""))/100))**(-datedif/days_in_quarter(Loop_Ending, dfdata))
        calc_df["Net VTD"] = 0
        calc_df["Net VTD"][Loop_Ending] = calc_df["TMV VTD"][Loop_Ending] - df["ER Corrective Contribution"][m]
        calc_df["Net to EOQ"] = calc_df["Net VTD"]* \
        (1+(float(calc_df["Returns"][Loop_Ending].replace("%",""))/100))**(datedif/days_in_quarter(Loop_Ending, dfdata))
        calc_df["Bring Returns Forward"] = 0
        calc_df["Bring Returns Forward"][Loop_Ending] = calc_df["Net to EOQ"][Loop_Ending]
        for i in range(Loop_Ending+1, len(calc_df)):
            calc_df["Bring Returns Forward"][i] = calc_df["Bring Returns Forward"][i-1] \
            *(1+(float(calc_df["Returns"][i].replace("%",""))/100))
        calc_df["At Present"] = 0
        calc_df["At Present"][len(calc_df)-1] = calc_df["Bring Returns Forward"][len(calc_df)-1] \
        *(1+(float(Savings_Rate.replace("%",""))/100))**(Final_Val_Days/365)
        answer = round(calc_df["At Present"][len(calc_df)-1],2)
        df["ER Investment Earnings"][m] = answer
        if VTD_Date > dfdata["Ending Dates"][len(dfdata)-1]:
            dfdata = dfdata.drop(dfdata.index[-1])
    df["Total Investment Earnings"] = df["EE Investment Earnings"] + df["ER Investment Earnings"]
    return calc_df
################################################################
def positive_procedure_EE_ER (df, Single_VTD):
    df_positive = df[:-1]
    df_positive = df_positive[df_positive["Total Investment Earnings"] > 0]
    numeric_columns = ["Name", "EE Corrective Contribution", "EE Investment Earnings","ER Corrective Contribution","ER Investment Earnings", "Total Corrective Contribution", "Total Investment Earnings"]
    df_positive = df_positive.groupby(['Name'], as_index=False)[numeric_columns].sum(numeric_only=True)
    df_positive['Name'] = df_positive['Name'].astype(str)
    df_positive['Year'] = df_positive['Name'].str.split().str[-1]
    df_positive['Year'] = pd.to_numeric(df_positive['Year'])
    df_positive = df_positive.sort_values(by=['Year', 'Name'])
    df_positive.drop('Year', axis=1, inplace=True)
    sum_row = df_positive.sum(numeric_only=True)
    sum_df = pd.DataFrame([sum_row], columns=df.columns)
    df_positive = pd.concat([df_positive, sum_df], ignore_index=True)
    df_positive = df_positive[["Name", "EE Corrective Contribution", "EE Investment Earnings","ER Corrective Contribution","ER Investment Earnings", "Total Corrective Contribution", "Total Investment Earnings"]]
    df_positive.loc[len(df_positive)-1, "Name"] = "Total"
    output_file_name_1 = f"Positive Earnings VTD {Single_VTD}.xlsx"
    df_positive.to_excel(output_file_name_1, index=False)
    print(f"Dataframe exported to {output_file_name_1}")
    return df_positive

def negative_procedure_EE_ER (df, Single_VTD):
    df_negative = df[:-1]
    df_negative = df_negative[df_negative["Total Investment Earnings"] < 0]
    numeric_columns = ["Name", "EE Corrective Contribution", "EE Investment Earnings","ER Corrective Contribution","ER Investment Earnings", "Total Corrective Contribution", "Total Investment Earnings"]
    df_negative = df_negative.groupby(['Name'], as_index=False)[numeric_columns].sum(numeric_only=True)
    df_negative['Name'] = df_negative['Name'].astype(str)
    df_negative['Year'] = df_negative['Name'].str.split().str[-1]
    df_negative['Year'] = pd.to_numeric(df_negative['Year'])
    df_negative = df_negative.sort_values(by=['Year', 'Name'])
    df_negative.drop('Year', axis=1, inplace=True)
    sum_row = df_negative.sum(numeric_only=True)
    sum_df = pd.DataFrame([sum_row], columns=df.columns)
    df_negative = pd.concat([df_negative, sum_df], ignore_index=True)
    df_negative = df_negative[["Name", "EE Corrective Contribution", "EE Investment Earnings","ER Corrective Contribution","ER Investment Earnings", "Total Corrective Contribution", "Total Investment Earnings"]]
    df_negative['"Recommended" Investment Earnings'] = 0
    df_negative.loc[len(df_negative)-1, "Name"] = "Total"
    output_file_name_2 = f"Negative Earnings VTD {Single_VTD}.xlsx"
    df_negative.to_excel(output_file_name_2, index=False)
    print(f"Dataframe exported to {output_file_name_2}")  
    return df_negative

def summary_procedure_EE_ER (df, Single_VTD, contains_negative):
    df_summary = df[["Name", "EE Corrective Contribution", "EE Investment Earnings","ER Corrective Contribution","ER Investment Earnings", "Total Corrective Contribution", "Total Investment Earnings"]]
    numeric_columns = ["EE Corrective Contribution", "EE Investment Earnings","ER Corrective Contribution","ER Investment Earnings", "Total Corrective Contribution", "Total Investment Earnings"]
    df_summary = df_summary.groupby(['Name'], as_index=False)[numeric_columns].sum(numeric_only=True)
    df_summary['Name'] = df_summary['Name'].astype(str)
    df_summary['Year'] = df_summary['Name'].str.split().str[-1]
    df_summary['Year'] = pd.to_numeric(df_summary['Year'])
    df_summary = df_summary.sort_values(by=['Year', 'Name'])
    df_summary.drop('Year', axis=1, inplace=True)
    sum_row = df_summary.sum(numeric_only=True)
    sum_df = pd.DataFrame([sum_row], columns=df.columns)
    df_summary = pd.concat([df_summary, sum_df], ignore_index=True)
    df_summary = df_summary[["Name", "EE Corrective Contribution", "EE Investment Earnings","ER Corrective Contribution","ER Investment Earnings", "Total Corrective Contribution", "Total Investment Earnings"]]
    if contains_negative == "yes":
        df_summary['"Recommended" Investment Earnings'] = 0
    df_summary.loc[len(df_summary)-1, "Name"] = "Total"
    output_file_name = f"Summary VTD {Single_VTD}.xlsx"
    df_summary.to_excel(output_file_name, index=False)
    print(f"Dataframe exported to {output_file_name}")
    return df_summary
#####################################################################
def letter_maker(name1, name2, PlanType, er, ernum, fcm, fcy, corrcontr, invearn, mostrecent, vtd, accrual, garate, negreturns, SingleorMultiple):
    months = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
    
    if name2[-1] in ["s", "z"]:
        name2 = name2 + "'"
    else:
        name2 = name2 + "'s"

    firstcontribution_month = months.index(fcm) + 1

    if firstcontribution_month < 4:
        firstcontribution_quarter = "first quarter of "
        fcq = 1
    elif firstcontribution_month < 7:
        firstcontribution_quarter = "second quarter of "
        fcq = 2
    elif firstcontribution_month < 10:
        firstcontribution_quarter = "third quarter of "
        fcq = 3
    else:
        firstcontribution_quarter = "fourth quarter of "
        fcq = 4

    mostrecentmth = mostrecent.month
    mry = mostrecent.year

    vtdcontribution_month = vtd.month

    if vtdcontribution_month < 4:
        vtdq = 1
    elif vtdcontribution_month < 7:
        vtdq = 2
    elif vtdcontribution_month < 10:
        vtdq = 3
    else:
        vtdq = 4

    if mostrecentmth == 11:
        mrm = "October and November"
    elif mostrecentmth == 10:
        mrm = "October"
    elif mostrecentmth == 8:
        mrm = "July and August"
    elif mostrecentmth == 7:
        mrm = "July"
    elif mostrecentmth == 5:
        mrm = "April and May"
    elif mostrecentmth == 4:
        mrm = "April"
    elif mostrecentmth == 2:
        mrm = "January and February"
    elif mostrecentmth == 1:
        mrm = "January"

    var1 = "As you requested, we have used the average rate of return method to calculate the additional investment earnings that should be applied to"
    var1m = "As you requested, we have used the average rate of return method to calculate the additional investment earnings that should be applied to the"
    var2 = " account."
    var2m = " accounts of certain participants."
    var3 = "Additional investment earnings of"
    var3m = "The following additional investment earnings should be deposited into these participants' "
    var4 = " should be deposited into participant"
    var5 = " account for the"
    var5m = " accounts for the corrective contributions "
    var6 = " of corrective contributions"
    var7 = " deposited for this participant effective"
    var7m = " deposited for these participants effective"

    if ((vtd.year == fcy) and (vtdq - fcq == 1)) or ((vtd.year - fcy == 1) and (fcq == 4 and vtdq == 1)):
        var8 = "To determine these investment earnings we first estimated the average earnings rates for the entire plan for the "
    else:
        var8 = "To determine these investment earnings we first estimated the average earnings rates for the entire plan for each quarter beginning with the "

    if (mostrecentmth % 3) == 0:
        var9 = "."
    elif ((mostrecentmth + 1) % 3) == 0:
        var9 = ", and for the months of " + mrm + " " + str(mry) + "."
    elif ((mostrecentmth + 2) % 3) == 0:
        var9 = ", and for the month of " + mrm + " " + str(mry) + "."

    var10 = ", the quarter in which the earliest corrective contribution applies"
    var10b = "  Each rate is estimated by assuming that contributions, benefit payments and all other plan financial activity occurred at the mid-point of the period."
    var11 = "  These average earnings rates were applied to the corrective contributions from the dates the contributions should have been made through "
    var12A = mostrecent.strftime("%B %d, %Y") + ", the date to which information exists that allows us to determine average plan earnings."
    var12B = "the deposit date of " + vtd.strftime("%B %d, %Y") + "."

    if vtd <= mostrecent:
        var12 = var12B
    else:
        var12 = var12A

    var13 = "  We used a conservative approach for applying the earnings rates to the corrective contributions to assure the maximum return for the participant."
    var14n = "  Specifically, we assume that contributions that should have been made in a month were made on the last day of the prior month.  "
    var14y = "  Specifically, in a quarter when the plan earnings rate is positive we assume that contributions that should have been made in a month were made on the last day of the prior month.  In a quarter when the plan earnings rate is negative we assume that contributions made in a month were made on the last day of that month.  "

    if negreturns == "No":
        var14 = var14n
    else:
        var14 = var14y

    var15 = "We then continued the accumulation of the contributions through the deposit date of "
    var16 = " at the current annual General Account interest rate"
    var17 = "The difference between the accumulations and the amount of the corrective contributions is the additional investment earnings as of the deposit date of "
    var18 = ".  We credited these additional earnings with accruals through "
    var19 = ", the date by which you should deposit the total amount of "

    par0 = "Dear ,"

    if SingleorMultiple == "Single":
        par1 = var1 + " " + name1 + " " + name2 + " " + PlanType + var2
        par2 = var3 + " " + invearn + var4 + " " + name2 + " " + PlanType + var5 + " " + corrcontr + var6 + " " + er + var7 + " " + vtd.strftime("%B %d, %Y") + "."
        var20 = " into the participant's account."
    else:
        par1 = var1m + " " + PlanType + var2m
        par2 = var3m + PlanType + var5m + er + var7m + " " + vtd.strftime("%B %d, %Y") + ":"
        var20 = " into the participants' accounts."

    tem3 = var8 + firstcontribution_quarter + str(fcy) + var10 + var9 + var10b + var11 + var12

    if vtd > mostrecent:
        tem4 = var15 + vtd.strftime("%B %d, %Y") + var16 + ".  " + var17 + vtd.strftime("%B %d, %Y") + var18 + accrual.strftime("%B %d, %Y") + var19 + invearn + var20
    else:
        tem4 = var17 + vtd.strftime("%B %d, %Y") + var18 + accrual.strftime("%B %d, %Y") + var19 + invearn + var20

    par3 = tem3 + var13 + var14 + tem4

    document = Document()
    document.add_paragraph(par0)
    document.add_paragraph(par1)
    document.add_paragraph(par2)
    document.add_paragraph(par3)
    
    output_file_name = f"{ernum} Investment Earnings {vtd}.docx"
    document.save(f"{output_file_name}")
    # print(f"Document exported to {output_file_name}")

